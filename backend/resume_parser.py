#!/usr/bin/env python3
"""
Import as:

import backend.resume_parser as barepa
"""

import re
import os
import logging
from typing import Dict, List, Optional, Any
import pandas as pd
import pdfplumber
import PyPDF2
import nltk
from dataclasses import dataclass, asdict
# Add docx support
from docx import Document

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ResumeData:
    """Data class to store parsed resume information"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None
    skills: List[str] = None
    education: List[Dict[str, str]] = None
    experience: List[Dict[str, str]] = None
    projects: List[Dict[str, str]] = None
    certifications: List[str] = None
    languages: List[str] = None
    raw_text: str = ""
    file_type: str = ""  # Add file type tracking
    
    def __post_init__(self):
        if self.skills is None:
            self.skills = []
        if self.education is None:
            self.education = []
        if self.experience is None:
            self.experience = []
        if self.projects is None:
            self.projects = []
        if self.certifications is None:
            self.certifications = []
        if self.languages is None:
            self.languages = []


class ResumeParser:
    """A comprehensive resume parser for PDF and DOCX files"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
        self.skills_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php',
            'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'matlab', 'sql',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django',
            'flask', 'spring', 'laravel', 'bootstrap', 'jquery', 'webpack',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'cassandra',
            'elasticsearch', 'dynamodb',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github',
            'gitlab', 'terraform', 'ansible', 'chef', 'puppet',
            
            # Data Science & ML
            'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
            'matplotlib', 'seaborn', 'plotly', 'jupyter', 'spark', 'hadoop',
            
            # Other Technologies
            'linux', 'unix', 'windows', 'macos', 'api', 'rest', 'graphql', 'microservices',
            'agile', 'scrum', 'kanban', 'jira', 'confluence'
        ]
        
        self.degree_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'associate', 'diploma',
            'b.s.', 'b.a.', 'm.s.', 'm.a.', 'b.tech', 'm.tech', 'mba', 'md'
        ]
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            full_text = '\n'.join(text)
            logger.info(f"Successfully extracted text from DOCX: {docx_path}")
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
            raise Exception(f"Could not extract text from DOCX: {docx_path}")
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using multiple methods for better accuracy"""
        text = ""
        
        # Method 1: Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            logger.info(f"Successfully extracted text using pdfplumber from {pdf_path}")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
            
            # Method 2: Fallback to PyPDF2
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                logger.info(f"Successfully extracted text using PyPDF2 from {pdf_path}")
            except Exception as e:
                logger.error(f"Both PDF extraction methods failed: {e}")
                raise Exception(f"Could not extract text from PDF: {pdf_path}")
        
        return text.strip()
    
    def extract_text_from_file(self, file_path: str) -> tuple[str, str]:
        """Extract text from supported file formats"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path), 'pdf'
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path), 'docx'
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {self.supported_formats}")
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from text"""
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
            'portfolio': None
        }
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone number extraction
        phone_patterns = [
            r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'\b\d{10}\b',
            r'\(\d{3}\)\s*\d{3}-\d{4}',
            r'\d{3}-\d{3}-\d{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                if isinstance(phones[0], tuple):
                    contact_info['phone'] = ''.join(phones[0])
                else:
                    contact_info['phone'] = phones[0]
                break
        
        # LinkedIn extraction
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact_info['linkedin'] = linkedin_matches[0]
        
        # GitHub extraction
        github_pattern = r'github\.com/[\w-]+'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact_info['github'] = github_matches[0]
        
        # Portfolio/Website extraction
        portfolio_pattern = r'https?://(?:www\.)?[\w.-]+\.[\w]{2,}'
        portfolio_matches = re.findall(portfolio_pattern, text)
        # Filter out common social media sites
        excluded_domains = ['linkedin.com', 'github.com', 'facebook.com', 'twitter.com']
        for match in portfolio_matches:
            if not any(domain in match for domain in excluded_domains):
                contact_info['portfolio'] = match
                break
        
        return contact_info
    
    def extract_name(self, text: str) -> Optional[str]:
        """Extract name from resume text"""
        lines = text.split('\n')
        
        # Try to get name from first few lines
        for line in lines[:5]:
            line = line.strip()
            if len(line.split()) >= 2 and len(line.split()) <= 4:
                # Check if line looks like a name (no numbers, not too long)
                if not re.search(r'\d', line) and len(line) < 50:
                    # Exclude common header words
                    exclude_words = ['resume', 'cv', 'curriculum', 'vitae', 'contact', 'phone', 'email']
                    if not any(word.lower() in line.lower() for word in exclude_words):
                        return line.strip()
        
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.skills_keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill.title())
        
        # Remove duplicates and sort
        return sorted(list(set(found_skills)))
    
    def extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        lines = text.split('\n')
        
        education_section = False
        current_edu = {}
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if we're in education section
            if any(keyword in line_lower for keyword in ['education', 'academic', 'qualification']):
                education_section = True
                continue
            
            # Stop if we hit another section
            if education_section and any(keyword in line_lower for keyword in ['experience', 'work', 'employment', 'project', 'skill']):
                if current_edu:
                    education.append(current_edu)
                break
            
            if education_section and line.strip():
                # Look for degree patterns
                for degree in self.degree_keywords:
                    if degree in line_lower:
                        if current_edu:
                            education.append(current_edu)
                        current_edu = {'degree': line.strip()}
                        break
                
                # Look for years
                year_match = re.search(r'(19|20)\d{2}', line)
                if year_match and 'degree' in current_edu:
                    current_edu['year'] = year_match.group()
                
                # Look for institution names (lines that don't contain degree keywords)
                if 'degree' in current_edu and 'institution' not in current_edu:
                    if not any(degree in line_lower for degree in self.degree_keywords):
                        if len(line.strip()) > 3:
                            current_edu['institution'] = line.strip()
        
        if current_edu:
            education.append(current_edu)
        
        return education
    
    def extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience information"""
        experience = []
        lines = text.split('\n')
        
        experience_section = False
        current_exp = {}
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if we're in experience section
            if any(keyword in line_lower for keyword in ['experience', 'work', 'employment', 'career']):
                experience_section = True
                continue
            
            # Stop if we hit another section
            if experience_section and any(keyword in line_lower for keyword in ['education', 'project', 'skill', 'certification']):
                if current_exp:
                    experience.append(current_exp)
                break
            
            if experience_section and line.strip():
                # Look for job titles and companies
                if len(line.strip()) > 3:
                    if not current_exp:
                        current_exp = {'title': line.strip()}
                    elif 'title' in current_exp and 'company' not in current_exp:
                        current_exp['company'] = line.strip()
                    elif 'title' in current_exp and 'company' in current_exp:
                        # Look for dates
                        date_match = re.search(r'(19|20)\d{2}.*?(19|20)\d{2}|present', line_lower)
                        if date_match:
                            current_exp['duration'] = line.strip()
                            experience.append(current_exp)
                            current_exp = {}
        
        if current_exp:
            experience.append(current_exp)
        
        return experience
    
    def parse_resume(self, file_path: str) -> ResumeData:
        """Main method to parse a resume file (PDF or DOCX)"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        # Check file format
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {self.supported_formats}")
        
        logger.info(f"Starting to parse resume: {file_path} (Format: {file_ext})")
        
        # Extract text from file
        text, file_type = self.extract_text_from_file(file_path)
        
        if not text:
            raise ValueError(f"No text could be extracted from the {file_type.upper()} file")
        
        # Create resume data object
        resume_data = ResumeData()
        resume_data.raw_text = text
        resume_data.file_type = file_type
        
        # Extract various information
        resume_data.name = self.extract_name(text)
        
        contact_info = self.extract_contact_info(text)
        resume_data.email = contact_info['email']
        resume_data.phone = contact_info['phone']
        resume_data.linkedin = contact_info['linkedin']
        resume_data.github = contact_info['github']
        resume_data.portfolio = contact_info['portfolio']
        
        resume_data.skills = self.extract_skills(text)
        resume_data.education = self.extract_education(text)
        resume_data.experience = self.extract_experience(text)
        
        logger.info(f"Resume parsing completed successfully (Format: {file_type})")
        return resume_data
    
    def save_to_json(self, resume_data: ResumeData, output_path: str):
        """Save parsed resume data to JSON file"""
        import json
        
        data_dict = asdict(resume_data)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Resume data saved to: {output_path}")
    
    def save_to_csv(self, resume_data: ResumeData, output_path: str):
        """Save parsed resume data to CSV file"""
        # Flatten the data for CSV format
        flattened_data = {
            'name': resume_data.name,
            'email': resume_data.email,
            'phone': resume_data.phone,
            'linkedin': resume_data.linkedin,
            'github': resume_data.github,
            'portfolio': resume_data.portfolio,
            'skills': ', '.join(resume_data.skills) if resume_data.skills else '',
            'education_count': len(resume_data.education),
            'experience_count': len(resume_data.experience),
            'file_type': resume_data.file_type
        }
        
        df = pd.DataFrame([flattened_data])
        df.to_csv(output_path, index=False)
        
        logger.info(f"Resume data saved to: {output_path}")


def main():
    """Example usage of the ResumeParser"""
    parser = ResumeParser()
    
    # Example usage for both PDF and DOCX
    file_paths = ["sample_resume.pdf", "sample_resume.docx"]  # Replace with your file paths
    
    for file_path in file_paths:
        if os.path.exists(file_path):
            try:
                # Parse the resume
                resume_data = parser.parse_resume(file_path)
                
                # Print extracted information
                print("\n" + "="*50)
                print(f"RESUME PARSING RESULTS - {resume_data.file_type.upper()}")
                print("="*50)
                
                print(f"Name: {resume_data.name}")
                print(f"Email: {resume_data.email}")
                print(f"Phone: {resume_data.phone}")
                print(f"LinkedIn: {resume_data.linkedin}")
                print(f"GitHub: {resume_data.github}")
                print(f"Portfolio: {resume_data.portfolio}")
                
                print(f"\nSkills ({len(resume_data.skills)}):")
                for skill in resume_data.skills[:10]:  # Show first 10 skills
                    print(f"  - {skill}")
                
                print(f"\nEducation ({len(resume_data.education)}):")
                for edu in resume_data.education:
                    print(f"  - {edu}")
                
                print(f"\nExperience ({len(resume_data.experience)}):")
                for exp in resume_data.experience:
                    print(f"  - {exp}")
                
                # Save results
                base_name = os.path.splitext(file_path)[0]
                parser.save_to_json(resume_data, f"{base_name}_parsed.json")
                parser.save_to_csv(resume_data, f"{base_name}_parsed.csv")
                
            except Exception as e:
                logger.error(f"Error parsing resume {file_path}: {e}")


if __name__ == "__main__":
    main()